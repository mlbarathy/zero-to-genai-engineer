"""Generate reproducible sample documents for the ingestion & chunking notebooks.

Creates, in this folder:
  - sample_report.pdf     text + a real table (the "parsing quality-loss" demo)
  - sample_page.html      headings, paragraphs, a list, and a table
  - sample_contract.docx  numbered contract clauses
  - sample_scanned.png    a memo rendered as an image (for OCR)

Run:  python3 _generate_sample_data.py
"""
from pathlib import Path

HERE = Path(__file__).resolve().parent


# ── 1. PDF with a table ──────────────────────────────────────────────────────
def make_pdf():
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                    TableStyle)

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(HERE / "sample_report.pdf"), pagesize=LETTER)
    story = []

    story.append(Paragraph("Northwind Logistics — Q2 2024 Operations Report", styles["Title"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "This report summarizes regional shipment performance for the second "
        "quarter of 2024. Overall shipment volume grew 8.4% quarter over quarter, "
        "driven primarily by the West and South regions. On-time delivery held "
        "above target in every region except the Northeast, where a warehouse "
        "migration temporarily reduced throughput.", styles["BodyText"]))
    story.append(Spacer(1, 12))
    story.append(Paragraph("Regional performance is broken out in the table below.",
                           styles["BodyText"]))
    story.append(Spacer(1, 12))

    data = [
        ["Region", "Shipments", "On-Time %", "Revenue ($M)"],
        ["Northwest", "128,400", "97.2%", "14.6"],
        ["West", "204,910", "98.1%", "23.9"],
        ["South", "176,220", "96.8%", "19.4"],
        ["Northeast", "92,180", "89.3%", "11.2"],
        ["Total", "601,710", "95.9%", "69.1"],
    ]
    table = Table(data, hAlign="LEFT")
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1a1a2e")),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
        ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
        ("BACKGROUND", (0, -1), (-1, -1), colors.HexColor("#faf6f0")),
        ("FONTNAME", (0, -1), (-1, -1), "Helvetica-Bold"),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
    ]))
    story.append(table)
    story.append(Spacer(1, 12))
    story.append(Paragraph(
        "Action items: (1) complete the Northeast warehouse migration by July 15; "
        "(2) reallocate two delivery routes from Northwest to West to match demand; "
        "(3) review carrier contracts ahead of the Q3 peak season.", styles["BodyText"]))
    doc.build(story)
    print("wrote sample_report.pdf")


# ── 2. HTML page ─────────────────────────────────────────────────────────────
def make_html():
    html = """<!DOCTYPE html>
<html lang="en">
<head><meta charset="utf-8"><title>Northwind Logistics — Support FAQ</title></head>
<body>
  <h1>Northwind Logistics — Support FAQ</h1>
  <p>Answers to common questions about shipping, tracking, and returns.</p>

  <h2>Shipping</h2>
  <p>Standard shipping takes 3–5 business days. Expedited shipping takes 1–2
  business days and is available in the West and Northwest regions.</p>
  <ul>
    <li>Free standard shipping on orders over $75.</li>
    <li>Expedited shipping is a flat $12.</li>
    <li>We do not ship to PO boxes.</li>
  </ul>

  <h2>Tracking</h2>
  <p>Every shipment includes a tracking number sent by email once the order
  leaves the warehouse. Tracking updates every 6 hours.</p>

  <h2>Returns</h2>
  <p>Returns are accepted within 30 days of delivery. Refunds are processed to
  the original payment method within 5 business days of receipt.</p>
  <table border="1">
    <tr><th>Item type</th><th>Return window</th><th>Restocking fee</th></tr>
    <tr><td>Standard goods</td><td>30 days</td><td>None</td></tr>
    <tr><td>Perishable</td><td>Not returnable</td><td>&mdash;</td></tr>
    <tr><td>Bulk / freight</td><td>14 days</td><td>10%</td></tr>
  </table>
</body>
</html>
"""
    (HERE / "sample_page.html").write_text(html, encoding="utf-8")
    print("wrote sample_page.html")


# ── 3. Word contract ─────────────────────────────────────────────────────────
def make_docx():
    from docx import Document

    doc = Document()
    doc.add_heading("Master Services Agreement", level=0)
    doc.add_paragraph(
        "This Master Services Agreement (\"Agreement\") is entered into between "
        "Northwind Logistics, Inc. (\"Provider\") and the Client, effective as of "
        "the date of last signature.")

    clauses = [
        ("1. Scope of Services",
         "Provider shall deliver freight logistics and warehousing services as "
         "described in each executed Statement of Work."),
        ("2. Term",
         "This Agreement remains in effect for an initial term of twelve (12) "
         "months and renews automatically for successive twelve-month terms unless "
         "either party gives sixty (60) days written notice of non-renewal."),
        ("3. Fees and Payment",
         "Client shall pay all undisputed invoices within thirty (30) days of "
         "receipt. Late payments accrue interest at 1.5% per month."),
        ("4. Termination",
         "Either party may terminate this Agreement for material breach if the "
         "breach remains uncured thirty (30) days after written notice. Provider "
         "may suspend services immediately for non-payment exceeding sixty (60) days."),
        ("5. Remote-Work Reimbursement (HR-2024-17 §7.3)",
         "Employees working remotely at least three (3) days per week are eligible "
         "for a home-internet stipend of $45 per month, reimbursed via the monthly "
         "expense report. Contractors are not eligible."),
        ("6. Limitation of Liability",
         "Neither party's aggregate liability shall exceed the fees paid under this "
         "Agreement in the twelve (12) months preceding the claim."),
    ]
    for title, body in clauses:
        doc.add_heading(title, level=1)
        doc.add_paragraph(body)

    doc.save(str(HERE / "sample_contract.docx"))
    print("wrote sample_contract.docx")


# ── 4. Scanned-style image (for OCR) ─────────────────────────────────────────
def make_scanned_png():
    from PIL import Image, ImageDraw, ImageFont

    lines = [
        "INTERNAL MEMO",
        "",
        "TO:   All warehouse staff",
        "FROM: Operations, Northwind Logistics",
        "DATE: 2024-06-14",
        "RE:   Incident INC-5521",
        "",
        "The claims-search service had a 22-minute outage on",
        "2024-06-14 caused by a Qdrant snapshot restore that",
        "exhausted disk. Fix: volume size increased and a disk-",
        "usage alert added at 80%. Owner: platform team.",
    ]
    W, H = 900, 520
    img = Image.new("RGB", (W, H), "white")
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Courier New.ttf", 24)
        title_font = ImageFont.truetype("/System/Library/Fonts/Supplemental/Arial Bold.ttf", 30)
    except OSError:
        font = ImageFont.load_default()
        title_font = font

    y = 30
    for i, line in enumerate(lines):
        draw.text((40, y), line, fill="black", font=title_font if i == 0 else font)
        y += 40 if i == 0 else 34
    # a light border to look like a scan
    draw.rectangle([5, 5, W - 6, H - 6], outline="#cccccc", width=2)
    img.save(str(HERE / "sample_scanned.png"))
    print("wrote sample_scanned.png")


# ── 5. Multimodal PDF — text + real embedded charts (Notebook 15) ───────────
def make_multimodal_pdf():
    """Northwind Logistics — Q3 2024 report, the sequel to sample_report.pdf's Q2
    report. The text narrates trends in *words only*; the exact figures live
    exclusively inside two embedded chart images. This is one half of the teaching
    point of Notebook 15: a text-only RAG pipeline can discuss the vibe of Q3
    but cannot answer "what was West region's exact Q3 revenue" or "what % of
    volume did FastFreight carry" — those numbers only exist as pixels.

    The other half: four *non-chart* illustrations (warehouse, delivery truck, a
    stack of boxes, a route map) with NO text or numbers anywhere in them —
    plain flat-color line-art, drawn with matplotlib shapes only. Charts test
    whether a system can read numbers off pixels. These test whether a system
    understands *what something looks like* — a genuinely different skill, and
    exactly the case unified multimodal embeddings (CLIP-style) are built for,
    while being notoriously weak at the chart-reading task. Six real, distinct
    images total is also what makes retrieval selectivity demonstrable: with
    only two near-identical chart images, every query matches "a chart" and
    both come back every time. With six visually distinct images, a query about
    a warehouse should retrieve the warehouse and nothing else."""
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import matplotlib.patches as mpatches
    import numpy as np
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.lib.units import inch
    from reportlab.platypus import Image as RLImage
    from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer

    tmp_paths = []

    def _save(fig, name):
        p = HERE / f"_tmp_{name}.png"
        fig.savefig(p, dpi=150, bbox_inches="tight")
        plt.close(fig)
        tmp_paths.append(p)
        return p

    # ── Chart 1 — a real bar chart; these numbers appear NOWHERE in the report text.
    regions = ["Northwest", "West", "South", "Northeast"]
    revenue = [15.8, 26.3, 21.7, 14.9]
    fig1, ax1 = plt.subplots(figsize=(6, 3.4))
    bars = ax1.bar(regions, revenue, color=["#4c72b0", "#dd8452", "#55a868", "#c44e52"])
    ax1.set_title("Q3 2024 Revenue by Region ($M)")
    ax1.set_ylabel("Revenue ($M)")
    for b, v in zip(bars, revenue):
        ax1.text(b.get_x() + b.get_width() / 2, v + 0.4, f"${v}M", ha="center", fontsize=10)
    ax1.set_ylim(0, max(revenue) + 4)
    fig1.tight_layout()
    chart1_path = _save(fig1, "revenue")

    # ── Chart 2 — a real pie chart; carrier mix, also nowhere in the text.
    carriers = ["FastFreight", "RegionalExpress", "NorthStar Logistics", "Other"]
    share = [38, 27, 21, 14]
    fig2, ax2 = plt.subplots(figsize=(5, 4))
    ax2.pie(share, labels=carriers, autopct="%1.0f%%", startangle=90,
            colors=["#4c72b0", "#dd8452", "#55a868", "#8172b2"])
    ax2.set_title("Q3 Shipment Volume by Carrier")
    fig2.tight_layout()
    chart2_path = _save(fig2, "carriers")

    # ── Icon 1 — warehouse building. Flat shapes only, zero text.
    fig3, ax3 = plt.subplots(figsize=(4, 4))
    ax3.set_xlim(0, 10); ax3.set_ylim(0, 10); ax3.set_aspect("equal"); ax3.axis("off")
    ax3.add_patch(mpatches.Rectangle((0, 0), 10, 0.4, color="#8a8a8a"))
    ax3.add_patch(mpatches.Rectangle((1.5, 0.4), 7, 4.5, color="#5b7fa6"))
    ax3.add_patch(mpatches.Polygon([[1.0, 4.9], [9.0, 4.9], [7.5, 6.8], [2.5, 6.8]], closed=True, color="#2e3f56"))
    ax3.add_patch(mpatches.Rectangle((2.2, 0.4), 2.2, 2.8, color="#c8ccd0"))
    for i in range(4):
        ax3.add_patch(mpatches.Rectangle((2.2, 0.4 + i * 0.7), 2.2, 0.05, color="#8a8f94"))
    ax3.add_patch(mpatches.Rectangle((5.8, 0.4), 1.0, 2.0, color="#3d3d3d"))
    for x in [3.0, 4.3, 5.6, 6.9]:
        ax3.add_patch(mpatches.Rectangle((x, 3.6), 0.7, 0.7, color="#dff0ff"))
    warehouse_path = _save(fig3, "warehouse")

    # ── Icon 2 — delivery truck, side view. Flat shapes only, zero text.
    fig4, ax4 = plt.subplots(figsize=(5, 3))
    ax4.set_xlim(0, 12); ax4.set_ylim(0, 6); ax4.set_aspect("equal"); ax4.axis("off")
    ax4.add_patch(mpatches.Rectangle((0, 0.6), 12, 0.15, color="#999999"))
    ax4.add_patch(mpatches.Rectangle((0.5, 1.3), 7, 2.8, color="#e8e8e8", ec="#333333", lw=1.5))
    ax4.add_patch(mpatches.FancyBboxPatch((7.5, 1.3), 3.2, 2.0, boxstyle="round,pad=0,rounding_size=0.3", color="#2f6fb0"))
    ax4.add_patch(mpatches.Polygon([[8.0, 2.6], [10.3, 2.6], [10.3, 3.3], [8.4, 3.3]], closed=True, color="#bde3ff"))
    for cx in [2.2, 5.5, 9.5]:
        ax4.add_patch(mpatches.Circle((cx, 0.75), 0.55, color="#222222"))
        ax4.add_patch(mpatches.Circle((cx, 0.75), 0.22, color="#888888"))
    truck_path = _save(fig4, "truck")

    # ── Icon 3 — stack of shipping boxes with tape lines. Flat shapes only, zero text.
    fig5, ax5 = plt.subplots(figsize=(4, 4))
    ax5.set_xlim(0, 10); ax5.set_ylim(0, 10); ax5.axis("off"); ax5.set_aspect("equal")
    for x, y, s, c in [(1.5, 0.5, 3.2, "#c98a4b"), (4.9, 0.5, 3.2, "#b97a3d"), (3.1, 3.5, 3.2, "#d99a5c")]:
        ax5.add_patch(mpatches.Rectangle((x, y), s, s, color=c, ec="#5a3d1f", lw=1.5))
        ax5.plot([x, x + s], [y + s, y], color="#5a3d1f", lw=2)
        ax5.plot([x, x + s], [y, y + s], color="#5a3d1f", lw=2)
    boxes_path = _save(fig5, "boxes")

    # ── Icon 4 — route map with start/end markers. Flat shapes only, zero text.
    fig6, ax6 = plt.subplots(figsize=(4, 4))
    ax6.set_xlim(0, 10); ax6.set_ylim(0, 10); ax6.axis("off"); ax6.set_aspect("equal")
    ax6.add_patch(mpatches.Rectangle((0, 0), 10, 10, color="#dff2e1"))
    xs = np.linspace(1, 9, 40)
    ys = 5 + 2 * np.sin(xs / 1.6)
    ax6.plot(xs, ys, color="#2f6fb0", lw=3, linestyle="--")
    ax6.add_patch(mpatches.Circle((1, ys[0]), 0.35, color="#2e7d32"))
    ax6.add_patch(mpatches.Circle((9, ys[-1]), 0.4, color="#c62828"))
    ax6.add_patch(mpatches.Polygon([[8.7, ys[-1] - 0.5], [9.3, ys[-1] - 0.5], [9, ys[-1] - 1.1]], closed=True, color="#c62828"))
    route_path = _save(fig6, "route")

    styles = getSampleStyleSheet()
    doc = SimpleDocTemplate(str(HERE / "sample_multimodal_report.pdf"), pagesize=LETTER)
    story = [
        Paragraph("Northwind Logistics — Q3 2024 Regional Performance Report", styles["Title"]),
        Spacer(1, 12),
        Paragraph(
            "Q3 2024 continued the growth trend from Q2. The West region maintained its "
            "lead and the South region closed the gap significantly. The Northeast region "
            "fully recovered from the warehouse migration issues that held it back last "
            "quarter. The chart below breaks out exact revenue contribution by region.",
            styles["BodyText"]),
        Spacer(1, 10),
        RLImage(str(chart1_path), width=5.2 * inch, height=2.9 * inch),
        Spacer(1, 14),
        Paragraph(
            "Carrier mix also shifted this quarter as we expanded our partnership with "
            "regional carriers to improve delivery times in dense urban corridors. The "
            "chart below shows the current shipment volume split across our carrier "
            "partners.", styles["BodyText"]),
        Spacer(1, 10),
        RLImage(str(chart2_path), width=4.2 * inch, height=3.4 * inch),
        PageBreak(),
        Paragraph("Operations Highlights", styles["Heading1"]),
        Spacer(1, 10),
        Paragraph(
            "This quarter we opened a new automated distribution warehouse in the West "
            "region to support increased volume. The facility below is now fully "
            "operational and handling overflow from our existing sites.", styles["BodyText"]),
        Spacer(1, 8),
        RLImage(str(warehouse_path), width=2.6 * inch, height=2.6 * inch),
        Spacer(1, 14),
        Paragraph(
            "Our delivery fleet expanded with additional box trucks to handle "
            "last-mile delivery in dense urban corridors, shown below.", styles["BodyText"]),
        Spacer(1, 8),
        RLImage(str(truck_path), width=3.2 * inch, height=1.9 * inch),
        Spacer(1, 14),
        Paragraph(
            "All fragile shipments are now double-boxed under our upgraded packaging "
            "protocol, illustrated below.", styles["BodyText"]),
        Spacer(1, 8),
        RLImage(str(boxes_path), width=2.4 * inch, height=2.4 * inch),
        Spacer(1, 14),
        Paragraph(
            "We also launched a new expedited route network connecting our regional "
            "hubs, mapped below.", styles["BodyText"]),
        Spacer(1, 8),
        RLImage(str(route_path), width=2.6 * inch, height=2.6 * inch),
        Spacer(1, 14),
        Paragraph(
            "Looking ahead to Q4, leadership expects the West and South regions to "
            "remain the primary growth drivers, with continued investment in regional "
            "carrier partnerships.", styles["BodyText"]),
    ]
    doc.build(story)
    for p in tmp_paths:
        p.unlink()
    print("wrote sample_multimodal_report.pdf (6 images: 2 charts + 4 non-text illustrations)")


if __name__ == "__main__":
    make_pdf()
    make_html()
    make_docx()
    make_scanned_png()
    make_multimodal_pdf()
    print("Done. Sample data in", HERE)
